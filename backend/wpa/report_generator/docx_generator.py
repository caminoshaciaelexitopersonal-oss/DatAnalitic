from docx import Document
from docx.shared import Inches
from io import BytesIO
from backend.core.state_store import StateStore

def create_docx_report(job_id: str, state_store: StateStore) -> bytes:
    """
    Creates a DOCX report in-memory using artifacts from the StateStore.
    """
    document = Document()
    document.add_heading(f'SADI Analysis Report: Job {job_id}', 0)

    # Add summary from manifest
    document.add_heading('Executive Summary', level=1)
    manifest = state_store.load_json_artifact(job_id, "manifest.json")
    if manifest:
        summary = manifest.get("summary", "No executive summary available.")
        document.add_paragraph(summary)
    else:
        document.add_paragraph("Manifest not found. No summary available.")

    # Add EDA results
    document.add_heading('Exploratory Data Analysis (EDA)', level=1)
    eda_summary = state_store.load_json_artifact(job_id, "eda/summary_statistics.json")
    if eda_summary:
        document.add_paragraph(f"Found {len(eda_summary)} columns.")
        # Add more detailed stats if needed
    else:
        document.add_paragraph("No EDA summary available.")

    # Add target detection results
    document.add_heading('Target Detection', level=1)
    target_info = state_store.load_json_artifact(job_id, "target.json")
    if target_info:
        document.add_paragraph(f"Decision Mode: {target_info.get('decision_mode')}")
        document.add_paragraph(f"Selected Target: {target_info.get('selected_target', 'None')}")
        document.add_paragraph(f"Explanation: {target_info.get('explanation')}")
    else:
        document.add_paragraph("No target detection information available.")

    # Add SHAP summary plot if it exists
    document.add_heading('Model Explainability (SHAP)', level=1)
    try:
        shap_plot_bytes = state_store.load_artifact_as_bytes(job_id, "explainability/shap_summary_bar.png")
        if shap_plot_bytes:
            image_stream = BytesIO(shap_plot_bytes)
            document.add_picture(image_stream, width=Inches(6.0))
        else:
            document.add_paragraph("SHAP summary plot not found.")
    except Exception:
        document.add_paragraph("Could not load SHAP summary plot.")

    # Save document to an in-memory stream
    doc_stream = BytesIO()
    document.save(doc_stream)
    doc_stream.seek(0)

    return doc_stream.getvalue()
